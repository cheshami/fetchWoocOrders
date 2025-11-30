import json
import pandas as pd
from docx import Document
from docx.shared import Pt
import os
import logging
from logging.handlers import RotatingFileHandler


def read_excel(excel_file):
    """Load the Excel file and filter rows based on status."""
    try:
        df = pd.read_excel(excel_file)
        df.columns = df.columns.str.strip()
        logging.info("Excel file loaded successfully.")

        if len(df.columns) != len(COLUMN_HEADERS):
            logging.error(f"Expected {len(COLUMN_HEADERS)} columns, but found {len(df.columns)}.")
            return None
        
        column_mapping = {old_name: new_name for old_name, new_name in zip(df.columns, COLUMN_HEADERS)}
        df.rename(columns=column_mapping, inplace=True)
        logging.info("Column names mapped successfully.")
        
        df = filter_processing_orders(df)
        return df
    except FileNotFoundError:
        logging.error(f"Excel file '{excel_file}' not found.")
        return None
    except ValueError as ve:
        logging.error(f"Value error while processing Excel file: {ve}")
        return None
    except Exception as e:
        logging.error(f"Error loading Excel file '{excel_file}': {e}")
        return None

def filter_processing_orders(df):
    """Filter rows where the 'status' column is 'در حال پردازش'."""
    if 'status' not in df.columns:
        logging.error("The 'status' column is missing from the DataFrame.")
        return pd.DataFrame()  # Return an empty DataFrame if 'status' is not found
    
    filtered_df = df[df['status'].str.strip().str.lower() == PROCESSING_STATUS]
    logging.info(f"Filtered DataFrame to {len(filtered_df)} rows with status '{PROCESSING_STATUS}'.")
    return filtered_df

def find_and_replace(doc, find_text, replace_text):
    """Find and replace text in the document."""
    replaced = replace_in_paragraphs(doc.paragraphs, find_text, replace_text)

    if not replaced:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if replace_in_paragraphs(cell.paragraphs, find_text, replace_text):
                        logging.info(f"Replaced '{find_text}' with '{replace_text}' in Document.")
                        return doc
    return doc

def replace_in_paragraphs(paragraphs, find_text, replace_text):
    """Replace text in specified paragraphs."""
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text, 1)
                set_run_style(run, find_text)
                return True
    return False

def set_run_style(run, find_text):
    """Set font style for the run."""
    if find_text == '__name__':
        run.font.name = config['font']['name']
    else:
        run.font.name = config['font']['address_text']
    run.font.size = Pt(config['font']['size1'])

def read_doc_file(doc_path):
    """Read the Word document."""
    if not os.path.exists(doc_path):
        logging.error(f"The file '{doc_path}' does not exist.")
        raise FileNotFoundError(f"The file '{doc_path}' does not exist.")
    
    return Document(doc_path)

def save_doc_file(doc_path, doc):
    """Save the modified Word document."""
    modified_doc_path = os.path.basename(ADDRESS_DOC_FILENAME)
    doc.save(modified_doc_path)
    logging.info(f"Modified document saved as '{modified_doc_path}'.")

def process_replacements(doc, df):
    """Process replacements for each row in the DataFrame."""
    for index, row in df.iterrows():
        try:
            doc = replace_placeholder(doc, '__name__', str(row['billing_name']))
            result = row['state_city'].split('،')
            if result[0].strip() == result[1].strip():
                row['state_city'] = result[0].strip()
            doc = replace_placeholder(doc, '__address__', f"{row['state_city']}، {row['address']}")

            # Process phone number
            phone_number = str(row['phone']).strip() if pd.notna(row['phone']) else ''
            if phone_number.endswith('.0'):
                phone_number = phone_number[:-2]  # Remove the '.0' part
            if phone_number.startswith('9') and len(phone_number) == 10:
                phone_number = '0' + phone_number  # Prepend '0' if it starts with '9'
            doc = replace_placeholder(doc, '__phone__', phone_number)

            # Process postcode
            postcode = str(row['postcode']).strip() if pd.notna(row['postcode']) else ''
            if postcode.endswith('.0'):
                postcode = postcode[:-2]  # Remove the '.0' part
            doc = replace_placeholder(doc, '__postcode__', postcode)
        except Exception as e:
            logging.error(f"Failed to process row {index}: {e}")
    return doc

def replace_placeholder(doc, placeholder, replacement):
    """Replace a placeholder in the document."""
    return find_and_replace(doc, placeholder, replacement)

if __name__ == "__main__":
    try:
        # Load configuration from config.json
        with open('config.json', 'r', encoding='utf-8') as config_file:
            config = json.load(config_file)
        
        EXCEL_FILE_PATH = config['EXCEL_FILE_PATH']
        DOC_TEMPLATE_PATH = config['DOC_TEMPLATE_PATH']
        ADD_LOG_FILE_PATH = config['ADD_LOG_FILE_PATH']
        ADDRESS_DOC_FILENAME = config['ADD_DOC_FILE_PATH']
        
        # Configure logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                RotatingFileHandler(ADD_LOG_FILE_PATH, encoding='utf-8', maxBytes=5 * 1024 * 1024, backupCount=2),
                logging.StreamHandler()
            ]
        )
        logging.info("Starting the document processing script.")

        if config['lang'] == "en":
            from mapping import ENGLISH_STATUS, ENGLISH_COLUMN_HEADERS
            PROCESSING_STATUS = ENGLISH_STATUS['processing']
            COLUMN_HEADERS = ENGLISH_COLUMN_HEADERS
        elif config['lang'] == "fa":
            from mapping import PERSIAN_STATUS, PERSIAN_COLUMN_HEADERS
            PROCESSING_STATUS = PERSIAN_STATUS['processing']
            COLUMN_HEADERS = PERSIAN_COLUMN_HEADERS
        
        doc = read_doc_file(DOC_TEMPLATE_PATH)
        df = read_excel(EXCEL_FILE_PATH)
        
        if df is not None and not df.empty:
            doc = process_replacements(doc, df)
            save_doc_file(DOC_TEMPLATE_PATH, doc)
        else:
            logging.warning("No valid data found in the Excel file.")
    except Exception as e:
        logging.error(f"An error occurred: {e}")